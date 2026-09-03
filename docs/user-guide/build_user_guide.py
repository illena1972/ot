from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT_DIR = Path(__file__).resolve().parent
OUTPUT_PATH = OUTPUT_DIR / "Руководство пользователя СпецОдежда v2.1.docx"

BLUE = "1D4ED8"
DARK_BLUE = "163B65"
LIGHT_BLUE = "EAF2FF"
PALE_GRAY = "F3F4F6"
MID_GRAY = "64748B"
DARK = "1F2937"
RED = "B42318"
GREEN = "177245"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_width = sum(widths_dxa)
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(table_width))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Страница ")
    set_run_font(run, size=9, color=MID_GRAY)
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Оглавление обновится при открытии документа"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, placeholder, fld_end])


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    keep_with_next(paragraph)
    return paragraph


def add_body(doc, text, bold_prefix=None):
    paragraph = doc.add_paragraph(style="Normal")
    if bold_prefix and text.startswith(bold_prefix):
        prefix_run = paragraph.add_run(bold_prefix)
        set_run_font(prefix_run, bold=True)
        body_run = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(body_run)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def create_numbering_instance(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    nsid = OxmlElement("w:nsid")
    nsid.set(qn("w:val"), f"{abstract_id + 1:08X}")
    abstract.append(nsid)
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    level.append(start)
    level.append(num_fmt)
    level.append(level_text)
    level.append(suffix)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    # A tiny per-sequence offset prevents Word from merging independent lists.
    sequence_indent = 540 + abstract_id
    tab.set(qn("w:pos"), str(sequence_indent))
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), str(sequence_indent))
    indent.set(qn("w:hanging"), "270")
    p_pr.append(tabs)
    p_pr.append(indent)
    level.append(p_pr)
    abstract.append(level)
    first_num_index = next(
        (
            index
            for index, child in enumerate(numbering)
            if child.tag == qn("w:num")
        ),
        len(numbering),
    )
    numbering.insert(first_num_index, abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return num_id


def get_paragraph_num_id(paragraph):
    p_pr = paragraph._p.pPr
    if p_pr is None or p_pr.numPr is None or p_pr.numPr.numId is None:
        return None
    return int(p_pr.numPr.numId.val)


def add_step(doc, text):
    previous_num_id = get_paragraph_num_id(doc.paragraphs[-1]) if doc.paragraphs else None
    num_id = previous_num_id if previous_num_id is not None else create_numbering_instance(doc)
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.paragraph_format.left_indent = Cm(0.95)
    paragraph.paragraph_format.first_line_indent = -Cm(0.45)
    paragraph.paragraph_format.space_after = Pt(4)
    num_pr = paragraph._p.get_or_add_pPr().get_or_add_numPr()
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_note(doc, title, text, kind="info"):
    palette = {
        "info": (LIGHT_BLUE, DARK_BLUE),
        "warning": ("FFF4E5", "8A4B00"),
        "danger": ("FDECEC", RED),
        "success": ("EAF7EF", GREEN),
    }
    fill, color = palette[kind]
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    label = paragraph.add_run(f"{title}. ")
    set_run_font(label, bold=True, color=color)
    run = paragraph.add_run(text)
    set_run_font(run, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_two_column_table(doc, rows, headers=("Элемент", "Назначение")):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = headers[0]
    table.cell(0, 1).text = headers[1]
    set_repeat_table_header(table.rows[0])
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "E8EEF5")
        for run in cell.paragraphs[0].runs:
            set_run_font(run, bold=True, color=DARK_BLUE)
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        for cell in cells:
            for run in cell.paragraphs[0].runs:
                set_run_font(run, size=10.5)
    set_table_geometry(table, [2700, 6660])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_three_column_table(doc, headers, rows, widths=(2600, 3200, 3560)):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.cell(0, idx).text = header
    set_repeat_table_header(table.rows[0])
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "E8EEF5")
        for run in cell.paragraphs[0].runs:
            set_run_font(run, bold=True, color=DARK_BLUE)
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cells[idx].text = value
            for run in cells[idx].paragraphs[0].runs:
                set_run_font(run, size=10.2)
    set_table_geometry(table, list(widths))
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)
    section.header_distance = Cm(1)
    section.footer_distance = Cm(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2

    heading_tokens = {
        1: (16, BLUE, 18, 8),
        2: (13, BLUE, 14, 7),
        3: (12, DARK_BLUE, 10, 5),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name, left, hanging in (
        ("List Bullet", Cm(0.95), Cm(0.45)),
        ("List Bullet 2", Cm(1.65), Cm(0.45)),
        ("List Number", Cm(0.95), Cm(0.45)),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = left
        style.paragraph_format.first_line_indent = -hanging
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.2

    settings = doc.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.text = "СпецОдежда | Руководство пользователя | Версия 2.1"
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header_paragraph.runs:
        set_run_font(run, size=9, color=MID_GRAY)

    add_page_number(section.footer.paragraphs[0])


def build_document():
    doc = Document()
    configure_document(doc)

    # Cover: editorial_cover with an A4 print-profile override.
    for _ in range(4):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("СИСТЕМА УЧЕТА СПЕЦОДЕЖДЫ И СИЗ")
    set_run_font(run, size=11, bold=True, color=BLUE)
    kicker.paragraph_format.space_after = Pt(18)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("СпецОдежда")
    set_run_font(run, size=30, bold=True, color=DARK_BLUE)
    title.paragraph_format.space_after = Pt(8)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Руководство пользователя")
    set_run_font(run, size=18, color=DARK)
    subtitle.paragraph_format.space_after = Pt(24)

    meta = doc.add_table(rows=3, cols=2)
    meta.style = "Table Grid"
    meta_rows = (
        ("Версия программы", "2.1"),
        ("Назначение", "Учет сотрудников, спецодежды, СИЗ, склада и сроков эксплуатации"),
        ("Редакция документа", "2 сентября 2026 г."),
    )
    for row_idx, values in enumerate(meta_rows):
        for col_idx, value in enumerate(values):
            meta.cell(row_idx, col_idx).text = value
            if col_idx == 0:
                set_cell_shading(meta.cell(row_idx, col_idx), PALE_GRAY)
            for run in meta.cell(row_idx, col_idx).paragraphs[0].runs:
                set_run_font(run, bold=(col_idx == 0), size=10.5)
    set_table_geometry(meta, [2700, 6660])

    doc.add_paragraph()
    add_note(
        doc,
        "Для пользователя",
        "Работайте только по адресу своей организации. Данные разных организаций разделены и не отображаются друг у друга.",
        "info",
    )

    doc.add_page_break()
    add_heading(doc, "Оглавление", 1)
    toc = doc.add_paragraph()
    add_toc(toc)
    add_note(
        doc,
        "Примечание",
        "Если номера страниц не появились автоматически, откройте документ в Microsoft Word, щелкните оглавление и выберите «Обновить поле».",
        "info",
    )

    doc.add_page_break()
    add_heading(doc, "1. Назначение системы", 1)
    add_body(
        doc,
        "Система «СпецОдежда» предназначена для учета сотрудников, выданной им спецодежды и средств индивидуальной защиты, складских остатков, сроков эксплуатации и потребности в заказе новых позиций.",
    )
    add_heading(doc, "Основные возможности", 2)
    for item in (
        "ведение справочников подразделений, служб, должностей и каталога одежды;",
        "ведение карточек сотрудников и их размеров;",
        "учет поступления и текущих остатков на складе;",
        "оформление выдачи сотруднику с проверкой доступного количества;",
        "контроль даты окончания срока эксплуатации;",
        "просмотр отчета по сотруднику;",
        "расчет потребности для заказа на выбранную дату;",
        "выгрузка отчета для заказа и личной карточки сотрудника в Excel.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "2. Вход и завершение работы", 1)
    add_heading(doc, "Вход", 2)
    add_step(doc, "Откройте в браузере адрес, предоставленный вашей организации, например https://rostok.app.bioclean.ru.")
    add_step(doc, "Введите свой логин в поле «Логин».")
    add_step(doc, "Введите пароль в поле «Пароль» и нажмите «Войти».")
    add_step(doc, "После входа проверьте название организации в правом верхнем углу.")
    add_note(
        doc,
        "Важно",
        "Если указано чужое или неверное название организации, не вводите данные и сообщите администратору.",
        "warning",
    )
    add_heading(doc, "Выход", 2)
    add_body(doc, "Нажмите кнопку со значком выхода в правом верхнем углу. Особенно важно выходить из системы на общем компьютере.")

    add_heading(doc, "3. Интерфейс и порядок первоначального заполнения", 1)
    add_two_column_table(
        doc,
        (
            ("Левая панель", "Переход между справочниками, складом, выдачей, отчетами и учетными карточками."),
            ("Верхняя панель", "Название системы, текущая организация, имя пользователя и кнопка выхода."),
            ("Рабочая область", "Таблицы, фильтры, формы и результаты выбранного раздела."),
            ("Карандаш", "Редактирование выбранной записи."),
            ("Корзина", "Удаление записи после дополнительного подтверждения."),
        ),
    )
    add_heading(doc, "Рекомендуемая последовательность для новой организации", 2)
    for item in (
        "Создайте подразделения.",
        "Создайте службы.",
        "Создайте должности.",
        "Заполните каталог одежды и СИЗ.",
        "Добавьте сотрудников.",
        "Внесите начальные остатки на склад.",
        "После этого оформляйте выдачи и формируйте отчеты.",
    ):
        add_step(doc, item)

    add_heading(doc, "4. Справочники", 1)
    add_heading(doc, "Подразделения, службы и должности", 2)
    add_body(doc, "Эти справочники используются при создании сотрудника. Порядок работы в них одинаков.")
    add_step(doc, "Откройте нужный раздел в группе «Справочники».")
    add_step(doc, "Нажмите «Добавить подразделение», «Добавить службу» или «Добавить должность».")
    add_step(doc, "Введите наименование и нажмите «Добавить».")
    add_step(doc, "Для изменения нажмите карандаш, исправьте наименование и сохраните изменения.")
    add_step(doc, "Для удаления нажмите корзину и подтвердите действие.")
    add_note(
        doc,
        "Проверка дублей",
        "Наименования сравниваются без учета регистра и лишних пробелов. Например, «Склад» и «склад» считаются одним наименованием.",
        "info",
    )
    add_note(
        doc,
        "Осторожно",
        "Не удаляйте запись справочника, если она используется сотрудниками. Сначала исправьте связанные данные.",
        "warning",
    )

    add_heading(doc, "Каталог одежды", 2)
    add_body(doc, "Каталог содержит наименования всех учитываемых видов спецодежды и СИЗ.")
    add_step(doc, "Откройте «Каталог одежды» и нажмите «Добавить одежду».")
    add_step(doc, "Введите наименование.")
    add_step(doc, "Выберите тип: «Верхняя одежда», «Обувь» или «Безразмерная».")
    add_step(doc, "Нажмите «Добавить».")
    add_body(doc, "Кнопки над таблицей позволяют показать все позиции или только выбранный тип. Дубли наименований определяются без учета регистра.")
    add_three_column_table(
        doc,
        ("Тип", "Какие параметры используются", "Примеры"),
        (
            ("Верхняя одежда", "Размер и рост", "Костюм, куртка, халат"),
            ("Обувь", "Только размер", "Ботинки, сапоги"),
            ("Безразмерная", "Размер и рост не указываются", "Каска, перчатки общего типа, очки"),
        ),
    )

    add_heading(doc, "5. Сотрудники", 1)
    add_heading(doc, "Добавление сотрудника", 2)
    add_step(doc, "Откройте раздел «Сотрудники» и нажмите «Добавить сотрудника».")
    add_step(doc, "Заполните фамилию, имя и отчество.")
    add_step(doc, "Выберите пол, подразделение, службу и должность.")
    add_step(doc, "При наличии данных укажите рост, размер одежды и размер обуви.")
    add_step(doc, "Нажмите «Добавить».")
    add_two_column_table(
        doc,
        (
            ("Обязательные поля", "Фамилия, имя, отчество, пол, подразделение, служба и должность."),
            ("Дополнительные поля", "Рост, размер одежды и размер обуви."),
            ("Поиск", "Введите часть фамилии в поле поиска."),
            ("Фильтр", "Выберите подразделение или оставьте «Все подразделения»."),
        ),
        headers=("Параметр", "Описание"),
    )
    add_heading(doc, "Редактирование и дубли", 2)
    add_body(doc, "Для редактирования нажмите карандаш в строке сотрудника, измените данные и нажмите «Сохранить изменения».")
    add_note(
        doc,
        "Проверка дублей",
        "Сотрудник считается дублем, если совпадают ФИО, подразделение, служба и должность. Регистр букв в ФИО не учитывается.",
        "info",
    )
    add_note(
        doc,
        "Перед удалением",
        "Проверьте отчет по сотруднику и наличие связанных выдач. Удаление выполняйте только при ошибочном создании записи.",
        "warning",
    )

    add_heading(doc, "6. Остатки на складе", 1)
    add_body(doc, "Раздел показывает текущие остатки по каждому наименованию, размеру и росту. Наименования выводятся в алфавитном порядке.")
    add_heading(doc, "Добавление остатка", 2)
    add_step(doc, "Откройте «Остатки на складе» и нажмите «Добавить».")
    add_step(doc, "Выберите наименование из каталога.")
    add_step(doc, "Укажите параметры позиции согласно ее типу.")
    add_step(doc, "Введите количество. Оно должно быть больше нуля.")
    add_step(doc, "При необходимости добавьте примечание и нажмите «Сохранить».")
    add_three_column_table(
        doc,
        ("Тип", "Размер", "Рост"),
        (
            ("Верхняя одежда", "Обязателен", "Обязателен"),
            ("Обувь", "Обязателен", "Не указывается"),
            ("Безразмерная", "Не указывается", "Не указывается"),
        ),
        widths=(3200, 3080, 3080),
    )
    add_heading(doc, "Изменение и удаление", 2)
    add_body(doc, "Нажмите карандаш, чтобы изменить количество или примечание. В поле «Количество» укажите новый фактический остаток, затем сохраните запись.")
    add_body(doc, "Корзина удаляет выбранную складскую запись после подтверждения. Применяйте удаление только для ошибочно заведенной строки.")
    add_note(
        doc,
        "Не смешивайте параметры",
        "Одинаковое наименование с разными размерами или ростом учитывается отдельными складскими строками.",
        "warning",
    )

    add_heading(doc, "7. Выдача спецодежды", 1)
    add_heading(doc, "Оформление выдачи", 2)
    add_step(doc, "Откройте раздел «Выдача одежды».")
    add_step(doc, "Найдите и выберите сотрудника. Поле поддерживает поиск по ФИО.")
    add_step(doc, "Укажите дату выдачи с клавиатуры или выберите ее в календаре.")
    add_step(doc, "При необходимости заполните общее примечание.")
    add_step(doc, "Нажмите «Добавить позицию».")
    add_step(doc, "Выберите наименование, укажите размер и рост по типу позиции, количество, срок эксплуатации в месяцах и примечание.")
    add_step(doc, "Проверьте строку «Доступно на складе» и нажмите «Добавить».")
    add_step(doc, "Повторите добавление для остальных позиций.")
    add_step(doc, "Проверьте итоговый список и нажмите «Оформить выдачу».")
    add_step(doc, "После сообщения «Выдача оформлена» нажмите «ОК».")
    add_note(
        doc,
        "Контроль остатка",
        "Система не позволяет выдать больше указанного сочетания наименования, размера и роста, чем доступно на складе.",
        "success",
    )
    add_note(
        doc,
        "До оформления",
        "Позицию можно удалить из подготовленного списка. После оформления исправления выполняются через «Отчет по выдаче».",
        "info",
    )

    add_heading(doc, "8. Отчет по выдаче", 1)
    add_body(doc, "Раздел показывает выданную сотрудникам спецодежду, даты и состояние каждой позиции.")
    add_step(doc, "Откройте «Отчет по выдаче».")
    add_step(doc, "Введите ФИО или выберите подразделение для поиска сотрудника.")
    add_step(doc, "Нажмите строку сотрудника.")
    add_step(doc, "Просмотрите наименование, количество, размер, рост, дату выдачи, срок окончания и статус.")
    add_two_column_table(
        doc,
        (
            ("Редактировать", "Для действующей позиции можно изменить количество, срок эксплуатации и примечание."),
            ("Удалить", "Удаляет выбранную запись выдачи после подтверждения. Перед подтверждением внимательно прочитайте сообщение системы."),
            ("Списать", "Доступно для просроченной позиции и удаляет ее из текущего перечня выданного."),
        ),
        headers=("Действие", "Результат"),
    )
    add_note(
        doc,
        "Учетные действия",
        "Редактирование, удаление и списание влияют на учет. Выполняйте их только на основании фактической операции или исправления ошибки.",
        "warning",
    )

    add_heading(doc, "9. Отчет для заказа", 1)
    add_body(doc, "Отчет рассчитывает потребность в замене спецодежды, срок эксплуатации которой закончится не позднее выбранной даты.")
    add_step(doc, "Откройте «Отчет для заказа».")
    add_step(doc, "В поле «Сформировать на дату» введите дату или выберите ее в календаре.")
    add_step(doc, "При необходимости выберите тип: все, верхняя одежда, обувь или безразмерная.")
    add_step(doc, "Просмотрите рассчитанные количества.")
    add_step(doc, "Нажмите на наименование, чтобы увидеть сотрудников, для которых требуется замена.")
    add_three_column_table(
        doc,
        ("Столбец", "Что показывает", "Расчет"),
        (
            ("Требуется заменить", "Количество выданных позиций со сроком окончания до выбранной даты включительно", "Сумма по наименованию, размеру и росту"),
            ("Есть на складе", "Текущий складской остаток той же позиции", "По совпадающим наименованию, размеру и росту"),
            ("К заказу", "Недостающее количество", "Требуется заменить минус остаток, но не меньше нуля"),
        ),
    )
    add_note(
        doc,
        "Пример",
        "Если требуется заменить 35 касок, а на складе есть 10, в столбце «К заказу» будет указано 25.",
        "info",
    )
    add_heading(doc, "Экспорт и печать", 2)
    add_step(doc, "Установите дату и фильтр, которые должны попасть в файл.")
    add_step(doc, "Нажмите «Экспорт в Excel». Будет загружен файл order_report.xlsx.")
    add_step(doc, "Откройте файл в Excel. Строки сотрудников первоначально свернуты.")
    add_step(doc, "Для печати краткого отчета оставьте детализацию свернутой.")
    add_step(doc, "Для печати с ФИО раскройте группы строк знаком «+» слева от таблицы.")
    add_step(doc, "В окне печати проверьте ориентацию «Альбомная», формат A4 и размещение по ширине на одной странице.")
    add_body(doc, "Выгрузка содержит на одном листе сводные строки и связанную с ними детализацию по сотрудникам.")

    add_heading(doc, "10. Учетные карточки", 1)
    add_body(doc, "Личная карточка формируется по выбранному сотруднику на основании его данных и истории выдачи.")
    add_step(doc, "Откройте раздел «Учетные карточки».")
    add_step(doc, "Выберите сотрудника в списке.")
    add_step(doc, "Проверьте отображаемые должность и подразделение.")
    add_step(doc, "Нажмите «Сформировать карточку».")
    add_step(doc, "Откройте загруженный файл Excel и проверьте данные перед печатью.")
    add_two_column_table(
        doc,
        (
            ("Сведения о сотруднике", "ФИО, подразделение, должность, пол, рост, размер одежды и обуви."),
            ("Выданные позиции", "Наименование, дата выдачи и количество по имеющимся данным учета."),
            ("Имя файла", "Формируется по шаблону «Карточка_СИЗ_Фамилия_Имя.xlsx»."),
        ),
        headers=("Раздел карточки", "Содержание"),
    )
    add_note(
        doc,
        "Перед формированием",
        "Если в карточке отсутствуют рост или размеры, сначала дополните запись сотрудника в справочнике «Сотрудники».",
        "info",
    )

    add_heading(doc, "11. Работа с датами", 1)
    add_body(doc, "Поля дат используются при выдаче и формировании отчета для заказа.")
    add_bullet(doc, "Дату можно выбрать в русскоязычном календаре.")
    add_bullet(doc, "Дату можно ввести с клавиатуры в формате ДД.ММ.ГГГГ.")
    add_bullet(doc, "Для очистки выбранной даты используйте кнопку «Очистить» в календаре.")
    add_bullet(doc, "После ввода проверьте день, месяц и год перед сохранением или экспортом.")
    add_note(doc, "Пример", "Дата 5 сентября 2026 года вводится как 05.09.2026.", "info")

    add_heading(doc, "12. Сообщения и типичные ситуации", 1)
    add_three_column_table(
        doc,
        ("Сообщение или ситуация", "Причина", "Что сделать"),
        (
            ("Запись с таким наименованием уже существует", "Дубль в справочнике, включая отличие только регистром", "Найдите существующую запись и при необходимости отредактируйте ее"),
            ("Сотрудник с такими ФИО... уже существует", "Совпали ФИО, подразделение, служба и должность", "Проверьте существующую карточку сотрудника"),
            ("Недостаточно на складе", "Запрошенное количество больше доступного остатка", "Проверьте размер и рост, уменьшите количество или сначала внесите поступление"),
            ("Заполните все данные", "Не выбран сотрудник, дата или отсутствуют позиции выдачи", "Заполните обязательные поля и повторите оформление"),
            ("Сотрудники не найдены", "Фильтр не соответствует данным либо справочник пуст", "Очистите поиск и фильтр или добавьте сотрудников"),
            ("Нет данных", "На выбранную дату отсутствуют позиции для отчета", "Проверьте дату и выбранный тип одежды"),
            ("Не удалось выполнить вход", "Неверный логин, пароль или адрес организации", "Проверьте адрес и раскладку клавиатуры, затем обратитесь к администратору"),
        ),
        widths=(3000, 3000, 3360),
    )

    add_heading(doc, "13. Правила безопасной работы", 1)
    for item in (
        "Не передавайте свой пароль другим сотрудникам.",
        "Всегда проверяйте название организации после входа.",
        "Не используйте удаление для обычного движения одежды; оно предназначено для исправления ошибочных записей.",
        "Перед выдачей проверяйте сотрудника, дату, размер, рост, количество и срок эксплуатации.",
        "Перед печатью отчета проверяйте выбранную дату и состояние раскрытия строк сотрудников.",
        "После работы на общем компьютере выполняйте выход из системы.",
        "При систематической ошибке запишите последовательность действий и передайте ее администратору.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "14. Краткая памятка", 1)
    add_two_column_table(
        doc,
        (
            ("Начать работу", "Заполнить справочники, каталог, сотрудников и склад."),
            ("Оформить выдачу", "Выбрать сотрудника и дату, добавить позиции, проверить остаток, оформить."),
            ("Проверить сотрудника", "Открыть «Отчет по выдаче», найти сотрудника, открыть его список."),
            ("Подготовить заказ", "Выбрать дату в «Отчете для заказа», проверить «К заказу», экспортировать Excel."),
            ("Сформировать карточку", "Открыть «Учетные карточки», выбрать сотрудника, скачать Excel."),
            ("Завершить работу", "Нажать кнопку выхода в правом верхнем углу."),
        ),
        headers=("Задача", "Краткий порядок"),
    )

    add_note(
        doc,
        "Версия документа",
        "Руководство описывает возможности программы «СпецОдежда» версии 2.1. Функции норм выдачи и разделение прав между ролями в эту версию не входят.",
        "info",
    )

    core_props = doc.core_properties
    core_props.title = "СпецОдежда. Руководство пользователя"
    core_props.subject = "Руководство пользователя системы учета спецодежды и СИЗ"
    core_props.author = "СпецОдежда"
    core_props.keywords = "спецодежда, СИЗ, склад, выдача, руководство пользователя"

    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
